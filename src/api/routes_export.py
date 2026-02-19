"""
导出 API：按 canvas 输出 Markdown / Word。

支持通过 cite_key_format 参数控制引用格式：
- numeric: [1], [2], [3]
- hash: [a3f7b2c91e04]
- author_date: [Smith2023]
"""

import io
import re

from fastapi import APIRouter, HTTPException, Response

from src.api.schemas import ExportRequest, ExportResponse
from src.collaboration.canvas.canvas_manager import get_canvas
from src.collaboration.export.formatter import export_canvas_markdown
from src.collaboration.memory.session_memory import get_session_store

router = APIRouter(prefix="/export", tags=["export"])


def _add_runs_from_html_node(
    paragraph,
    node,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    monospace: bool = False,
) -> None:
    """将 HTML 行内节点写入到 docx paragraph（基础格式）。"""
    from bs4 import NavigableString, Tag

    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if monospace:
                run.font.name = "Courier New"
        return
    if not isinstance(node, Tag):
        return

    tag = (node.name or "").lower()
    if tag in {"br"}:
        paragraph.add_run("\n")
        return
    next_bold = bold or tag in {"strong", "b"}
    next_italic = italic or tag in {"em", "i"}
    next_underline = underline or tag in {"u"}
    next_monospace = monospace or tag in {"code"}
    for child in node.contents:
        _add_runs_from_html_node(
            paragraph,
            child,
            bold=next_bold,
            italic=next_italic,
            underline=next_underline,
            monospace=next_monospace,
        )


def _markdown_to_docx_bytes(markdown_content: str) -> bytes:
    """将 Markdown 转换为基础 docx 字节流。"""
    try:
        import markdown2
        from bs4 import BeautifulSoup
        from docx import Document
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"docx export dependency missing: {e}")

    html = markdown2.markdown(markdown_content or "", extras=["fenced-code-blocks", "tables"])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    root = soup.body or soup
    for elem in root.children:
        tag = getattr(elem, "name", None)
        if not tag:
            text = str(elem).strip()
            if text:
                doc.add_paragraph(text)
            continue

        tag = tag.lower()
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = max(0, min(int(tag[1]) - 1, 5))
            p = doc.add_heading(level=level)
            for child in elem.contents:
                _add_runs_from_html_node(p, child)
        elif tag == "p":
            p = doc.add_paragraph()
            for child in elem.contents:
                _add_runs_from_html_node(p, child)
        elif tag in {"ul", "ol"}:
            list_style = "List Number" if tag == "ol" else "List Bullet"
            for li in elem.find_all("li", recursive=False):
                p = doc.add_paragraph(style=list_style)
                for child in li.contents:
                    _add_runs_from_html_node(p, child)
        elif tag == "pre":
            doc.add_paragraph(elem.get_text())
        else:
            text = elem.get_text(" ", strip=True)
            if text:
                doc.add_paragraph(text)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _safe_filename(name: str) -> str:
    s = re.sub(r"[^A-Za-z0-9._-]+", "_", (name or "").strip()).strip("._")
    return s or "research_draft"


@router.post("", response_model=ExportResponse)
def export_canvas(body: ExportRequest) -> ExportResponse | Response:
    canvas_id = body.canvas_id or ""
    if not canvas_id and body.session_id:
        meta = get_session_store().get_session_meta(body.session_id)
        if meta:
            canvas_id = meta.get("canvas_id") or ""
    if not canvas_id:
        raise HTTPException(status_code=400, detail="canvas_id or session_id is required")

    c = get_canvas(canvas_id)
    if c is None:
        raise HTTPException(status_code=404, detail="canvas not found")

    fmt = (body.format or "markdown").lower()
    if fmt not in {"markdown", "docx"}:
        raise HTTPException(status_code=400, detail="only markdown and docx format are supported")

    # 支持 cite_key_format 参数覆盖配置
    cite_key_format = body.cite_key_format
    if cite_key_format and cite_key_format not in ("numeric", "hash", "author_date"):
        raise HTTPException(
            status_code=400,
            detail=f"invalid cite_key_format: {cite_key_format}, must be one of: numeric, hash, author_date"
        )

    content = export_canvas_markdown(c, cite_key_format=cite_key_format)
    if fmt == "docx":
        docx_bytes = _markdown_to_docx_bytes(content)
        filename = _safe_filename(c.working_title or c.topic or canvas_id) + ".docx"
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    return ExportResponse(format="markdown", content=content, canvas_id=canvas_id, session_id=body.session_id or "")
